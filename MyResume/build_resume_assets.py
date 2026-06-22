from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
RESUME_DIR = Path(__file__).resolve().parent
SOURCE = RESUME_DIR / "resume.json"
OUTPUT_DIR = RESUME_DIR / "output"
PUBLIC_DIR = ROOT / "public"
PDF_OUT = OUTPUT_DIR / "Cai-You-David-Resume.pdf"
DOCX_OUT = OUTPUT_DIR / "Cai-You-David-Resume.docx"
README_OUT = OUTPUT_DIR / "README.resume.md"
SITE_PDF = PUBLIC_DIR / "resume.pdf"
SITE_README = ROOT / "README.md"

BLUE = colors.HexColor("#1F4D78")
MID_BLUE = colors.HexColor("#2E74B5")
MUTED = colors.HexColor("#595959")
INK = colors.HexColor("#1F1F1F")
LIGHT = colors.HexColor("#E6ECF2")


def load_resume() -> dict[str, Any]:
    return json.loads(SOURCE.read_text(encoding="utf-8"))


def esc(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def markdown_resume(data: dict[str, Any]) -> str:
    lines = [
        f"# {data['name']}",
        "",
        f"**{data['headline']}**",
        "",
        f"{data['location']}  ",
        f"Email: [{data['email']}](mailto:{data['email']})  ",
        f"LinkedIn: [linkedin.com/in/caiyou]({data['linkedin']})",
        "",
        "## Summary",
        "",
    ]
    for paragraph in data["summary"]:
        lines += [paragraph, ""]

    lines += ["## Top Skills", ""]
    lines += [f"- {skill}" for skill in data["skills"]]
    lines += ["", "## Experience", ""]

    for item in data["experience"]:
        lines += [
            f"### {item['company']}",
            "",
            f"**{item['role']}**  ",
            f"{item['period']} · {item['location']}",
            "",
        ]
        lines += [f"- {bullet}" for bullet in item["bullets"]]
        lines.append("")

    lines += ["## Education", ""]
    for item in data["education"]:
        lines += [
            f"### {item['school']}",
            "",
            f"**{item['degree']}**  ",
            item["period"],
            "",
        ]

    lines += [
        "## About This Repository",
        "",
        "This repository powers my personal website and keeps my resume assets versioned in GitHub. The PDF resume is available at `/resume.pdf` after deployment.",
        "",
    ]
    return "\n".join(lines)


def build_pdf(data: dict[str, Any]) -> None:
    base = getSampleStyleSheet()
    styles = {
        "name": ParagraphStyle(
            "name",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=17,
            leading=19,
            textColor=BLUE,
            alignment=TA_CENTER,
            spaceAfter=2,
        ),
        "subtitle": ParagraphStyle(
            "subtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9.6,
            leading=10.8,
            textColor=INK,
            alignment=TA_CENTER,
            spaceAfter=2,
        ),
        "contact": ParagraphStyle(
            "contact",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=9,
            textColor=MUTED,
            alignment=TA_CENTER,
            spaceAfter=5,
        ),
        "h1": ParagraphStyle(
            "h1",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10.2,
            leading=11,
            textColor=MID_BLUE,
        ),
        "body": ParagraphStyle(
            "body",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.1,
            leading=9.0,
            textColor=INK,
            spaceAfter=3,
        ),
        "role": ParagraphStyle(
            "role",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=8.7,
            leading=9.4,
            textColor=INK,
            spaceAfter=0,
        ),
        "role_meta": ParagraphStyle(
            "role_meta",
            parent=base["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=8.1,
            leading=8.8,
            textColor=BLUE,
            spaceAfter=1,
        ),
        "bullet": ParagraphStyle(
            "bullet",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=7.8,
            leading=8.5,
            textColor=INK,
            leftIndent=10,
            firstLineIndent=-6,
            spaceAfter=1.1,
        ),
        "small": ParagraphStyle(
            "small",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=7.8,
            leading=8.5,
            textColor=INK,
            spaceAfter=1,
        ),
    }

    def section(title: str) -> list[Any]:
        return [
            Spacer(1, 5),
            Paragraph(esc(title.upper()), styles["h1"]),
            HRFlowable(width="100%", thickness=0.45, color=LIGHT, spaceBefore=1, spaceAfter=3),
        ]

    story: list[Any] = [
        Paragraph(esc(data["name"]), styles["name"]),
        Paragraph(esc(data["headline"]), styles["subtitle"]),
        Paragraph(
            esc(f"{data['location']} | {data['email']} | linkedin.com/in/caiyou"),
            styles["contact"],
        ),
    ]

    story += section("Summary")
    story.append(Paragraph(esc(" ".join(data["summary"])), styles["body"]))

    story += section("Top Skills")
    story.append(Paragraph(esc(", ".join(data["skills"])), styles["body"]))

    story += section("Experience")
    for item in data["experience"]:
        story.append(Paragraph(esc(item["company"]), styles["role"]))
        story.append(
            Paragraph(
                esc(f"{item['role']} | {item['period']} | {item['location']}"),
                styles["role_meta"],
            )
        )
        for bullet in item["bullets"]:
            story.append(Paragraph(f"- {esc(bullet)}", styles["bullet"]))

    story += section("Education")
    for item in data["education"]:
        story.append(
            Paragraph(
                esc(f"{item['school']} - {item['degree']} | {item['period']}"),
                styles["small"],
            )
        )

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        leftMargin=0.45 * inch,
        rightMargin=0.45 * inch,
        topMargin=0.32 * inch,
        bottomMargin=0.32 * inch,
    )
    doc.build(story)


def build_docx(data: dict[str, Any]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(data["name"])
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(data["headline"])
    subtitle_run.font.size = Pt(10)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run(f"{data['location']} | {data['email']} | linkedin.com/in/caiyou")
    contact_run.font.size = Pt(8.5)
    contact_run.font.color.rgb = RGBColor(89, 89, 89)

    def heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(46, 116, 181)

    heading("Summary")
    p = doc.add_paragraph(" ".join(data["summary"]))
    p.paragraph_format.line_spacing = 1.05

    heading("Top Skills")
    doc.add_paragraph(", ".join(data["skills"]))

    heading("Experience")
    for item in data["experience"]:
        p = doc.add_paragraph()
        company = p.add_run(item["company"])
        company.bold = True
        company.font.size = Pt(9.3)
        meta = p.add_run(f" | {item['role']} | {item['period']} | {item['location']}")
        meta.italic = True
        meta.font.size = Pt(8.8)
        for bullet in item["bullets"]:
            b = doc.add_paragraph(style="List Bullet")
            b.paragraph_format.space_after = Pt(1)
            b.add_run(bullet).font.size = Pt(8.4)

    heading("Education")
    for item in data["education"]:
        doc.add_paragraph(f"{item['school']} - {item['degree']} | {item['period']}")

    doc.save(DOCX_OUT)


def main() -> None:
    data = load_resume()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)

    markdown = markdown_resume(data)
    README_OUT.write_text(markdown, encoding="utf-8")
    SITE_README.write_text(markdown, encoding="utf-8")

    build_pdf(data)
    build_docx(data)
    shutil.copyfile(PDF_OUT, SITE_PDF)

    print(f"Updated {README_OUT.relative_to(ROOT)}")
    print(f"Updated {SITE_README.relative_to(ROOT)}")
    print(f"Updated {PDF_OUT.relative_to(ROOT)}")
    print(f"Updated {DOCX_OUT.relative_to(ROOT)}")
    print(f"Updated {SITE_PDF.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
